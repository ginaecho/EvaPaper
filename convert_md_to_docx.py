#!/usr/bin/env python3
"""Convert markdown to docx using python-docx."""
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def md_to_docx(md_path, docx_path):
    with open(md_path, 'r') as f:
        lines = f.readlines()
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')
        
        # Title (h1)
        if line.startswith('# ') and not line.startswith('## '):
            p = doc.add_heading(line[2:], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        
        # H2
        if line.startswith('## '):
            doc.add_heading(line[3:], level=1)
            i += 1
            continue
        
        # H3
        if line.startswith('### '):
            doc.add_heading(line[4:], level=2)
            i += 1
            continue
        
        # H4
        if line.startswith('#### '):
            doc.add_heading(line[5:], level=3)
            i += 1
            continue
        
        # Horizontal rule
        if line.strip() == '---':
            doc.add_paragraph()
            i += 1
            continue
        
        # Blockquote
        if line.startswith('> '):
            quote_text = line[2:]
            # Collect multi-line blockquotes
            j = i + 1
            while j < len(lines) and lines[j].startswith('> '):
                quote_text += '\n' + lines[j][2:].rstrip('\n')
                j += 1
            p = doc.add_paragraph(quote_text)
            p.style = 'Intense Quote'
            i = j
            continue
        
        # Bullet list
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            p = doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue
        
        # Numbered list (simple)
        if line.strip() and line.strip()[0].isdigit() and '. ' in line.strip()[:4]:
            text = line.strip().split('. ', 1)[1]
            p = doc.add_paragraph(text, style='List Number')
            i += 1
            continue
        
        # Bold text handling
        if line.strip():
            p = doc.add_paragraph()
            parts = line.split('**')
            for idx, part in enumerate(parts):
                run = p.add_run(part)
                if idx % 2 == 1:  # Odd indices are bold
                    run.bold = True
            i += 1
            continue
        
        # Empty line
        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue
        
        # Regular paragraph
        doc.add_paragraph(line)
        i += 1
    
    doc.save(docx_path)
    print(f"Saved: {docx_path}")

if __name__ == '__main__':
    md_to_docx('/root/.openclaw/workspace/AI_Agent_Governance_Three_Layer_Stack_and_Papers.md',
               '/root/.openclaw/workspace/AI_Agent_Governance_Three_Layer_Stack_and_Papers.docx')
