from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent

INPUTS = [
    ROOT / "AI_Agent_Governance_Three_Layer_Stack_and_Papers.md",
    ROOT / "Agent_Governance_Three_Questions_Synthesis.md",
]

TITLE_COLOR = RGBColor(28, 49, 88)
SUBTITLE_COLOR = RGBColor(74, 86, 107)
HEADING_COLORS = {
    1: RGBColor(35, 64, 118),
    2: RGBColor(45, 90, 153),
    3: RGBColor(58, 116, 183),
    4: RGBColor(68, 134, 204),
}
BODY_COLOR = RGBColor(34, 34, 34)
MUTED_COLOR = RGBColor(92, 92, 92)
CALLOUT_FILL = "EAF2FF"


@dataclass
class Block:
    kind: str
    level: int = 0
    text: str = ""
    items: list[str] | None = None
    rows: list[list[str]] | None = None


def clean_inline(text: str) -> str:
    text = text.strip()
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    return text.strip()


def parse_markdown(path: Path) -> list[Block]:
    lines = path.read_text(encoding="utf-8").splitlines()
    blocks: list[Block] = []
    i = 0
    while i < len(lines):
        raw = lines[i].rstrip()
        stripped = raw.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = clean_inline(stripped[level:].strip())
            blocks.append(Block(kind="heading", level=level, text=text))
            i += 1
            continue

        if stripped.startswith(">"):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_lines.append(clean_inline(lines[i].strip().lstrip(">").strip()))
                i += 1
            blocks.append(Block(kind="quote", text=" ".join(q for q in quote_lines if q)))
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row = [clean_inline(cell) for cell in lines[i].strip().strip("|").split("|")]
                if not all(set(cell) <= {"-", ":"} for cell in row):
                    table_lines.append(row)
                i += 1
            if table_lines:
                blocks.append(Block(kind="table", rows=table_lines))
            continue

        if re.match(r"^([-*]|\d+\.)\s+", stripped):
            items = []
            while i < len(lines):
                candidate = lines[i].strip()
                if not re.match(r"^([-*]|\d+\.)\s+", candidate):
                    break
                item = re.sub(r"^([-*]|\d+\.)\s+", "", candidate)
                items.append(clean_inline(item))
                i += 1
            blocks.append(Block(kind="list", items=items))
            continue

        para_lines = [stripped]
        i += 1
        while i < len(lines):
            candidate = lines[i].strip()
            if (
                not candidate
                or candidate.startswith("#")
                or candidate.startswith(">")
                or candidate.startswith("|")
                or re.match(r"^([-*]|\d+\.)\s+", candidate)
            ):
                break
            para_lines.append(candidate)
            i += 1
        blocks.append(Block(kind="paragraph", text=clean_inline(" ".join(para_lines))))

    return blocks


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_title_page(doc: Document, title: str, source_name: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = TITLE_COLOR
    run.font.name = "Aptos Display"

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Generated from {source_name}")
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = SUBTITLE_COLOR
    run.font.name = "Aptos"

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(18)


def render_doc(source: Path) -> Path:
    blocks = parse_markdown(source)
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    style = doc.styles["Normal"]
    style.font.name = "Aptos"
    style.font.size = Pt(10.5)
    style.font.color.rgb = BODY_COLOR

    title = next((b.text for b in blocks if b.kind == "heading" and b.level == 1), source.stem)
    add_title_page(doc, title, source.name)

    for block in blocks:
        if block.kind == "heading":
            level = min(max(block.level, 1), 4)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(block.text)
            run.bold = True
            run.font.name = "Aptos Display"
            run.font.color.rgb = HEADING_COLORS.get(level, TITLE_COLOR)
            run.font.size = {1: Pt(20), 2: Pt(16), 3: Pt(13), 4: Pt(11.5)}[level]
            continue

        if block.kind == "paragraph":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.add_run(block.text)
            continue

        if block.kind == "quote":
            table = doc.add_table(rows=1, cols=1)
            table.autofit = True
            cell = table.rows[0].cells[0]
            set_cell_shading(cell, CALLOUT_FILL)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(block.text)
            run.italic = True
            run.font.color.rgb = SUBTITLE_COLOR
            run.font.size = Pt(10.5)
            doc.add_paragraph()
            continue

        if block.kind == "list" and block.items:
            for item in block.items:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(2)
                p.add_run(item)
            continue

        if block.kind == "table" and block.rows:
            cols = max(len(row) for row in block.rows)
            table = doc.add_table(rows=len(block.rows), cols=cols)
            table.style = "Table Grid"
            for r, row in enumerate(block.rows):
                for c in range(cols):
                    text = row[c] if c < len(row) else ""
                    cell = table.rows[r].cells[c]
                    cell.text = text
                    if r == 0:
                        set_cell_shading(cell, "D9E6FF")
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
            doc.add_paragraph()

    output = source.with_suffix(".docx")
    doc.save(output)
    return output


def main() -> None:
    generated = [render_doc(path) for path in INPUTS]
    for path in generated:
        print(f"DOCX saved: {path}")


if __name__ == "__main__":
    main()
